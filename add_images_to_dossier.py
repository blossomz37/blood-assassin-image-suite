#!/usr/bin/env python3
"""
Script to add images to Blood_Assassin_Character_Dossier.docx
Automatically inserts character portraits and scene images with captions
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)

# Image mapping with titles and suggested placement
IMAGE_MAP = {
    "01_elara_nightshade_portrait.png": {
        "title": "Elara Nightshade - Portrait",
        "section": "Elara",
        "width": 4.0
    },
    "02_queen_lysandria_portrait.png": {
        "title": "Queen Lysandria - Portrait",
        "section": "Lysandria",
        "width": 4.0
    },
    "03_seraphiel_portrait.png": {
        "title": "Seraphiel - Portrait",
        "section": "Seraphiel",
        "width": 4.0
    },
    "04_prophecy_chamber.png": {
        "title": "The Prophecy Chamber",
        "section": "locations",
        "width": 5.5
    },
    "05_crimson_court_throne_room.png": {
        "title": "Crimson Court Throne Room",
        "section": "locations",
        "width": 5.5
    },
    "06_elara_vs_lysandria.png": {
        "title": "Elara vs Lysandria - Confrontation",
        "section": "scenes",
        "width": 5.5
    },
    "07_nightbringer_manifestation.png": {
        "title": "The Nightbringer Manifestation",
        "section": "scenes",
        "width": 5.5
    },
    "08_rebel_encampment.png": {
        "title": "Rebel Encampment",
        "section": "locations",
        "width": 5.5
    },
    "09_blood_moon_fortress.png": {
        "title": "Blood Moon Fortress",
        "section": "locations",
        "width": 5.5
    },
    "10_twilight_accord_signing.png": {
        "title": "The Twilight Accord Signing",
        "section": "scenes",
        "width": 5.5
    },
    "11_elara_midnight_hunt.png": {
        "title": "Elara - Midnight Hunt",
        "section": "Elara",
        "width": 5.5
    },
    "12_seraphiel_celestial_powers.png": {
        "title": "Seraphiel - Celestial Powers",
        "section": "Seraphiel",
        "width": 5.5
    }
}


def find_paragraph_by_keyword(doc, keywords):
    """Find a paragraph containing any of the keywords (case-insensitive)"""
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                return i
    return None


def insert_image_after_paragraph(doc, para_index, image_path, title, width_inches):
    """Insert an image after a specific paragraph with a caption"""
    if not os.path.exists(image_path):
        print(f"WARNING: Image not found: {image_path}")
        return False
    
    # Insert image paragraph after the target paragraph
    # We need to add it after para_index
    # Get the element after which we want to insert
    target_para = doc.paragraphs[para_index]
    
    # Add a new paragraph for the image
    new_para = target_para.insert_paragraph_before("")
    
    # Actually, insert_paragraph_before won't work as expected
    # We need to work with the underlying XML
    # Let's use a simpler approach: find the insertion point and add there
    
    # Better approach: add to end and note we'll need manual reordering
    # Or: we can insert in order from top to bottom
    
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    # Center align the image
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add caption paragraph
    caption_para = target_para.insert_paragraph_before("")
    caption_para.add_run(title).font.size = Pt(11)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.runs[0].font.italic = True
    
    # Add spacing after
    spacing_para = target_para.insert_paragraph_before("")
    
    return True


def add_images_smart_placement(doc, images_dir):
    """Add images with smart placement based on document structure"""
    
    added_images = []
    
    # Process images in order
    for img_filename in sorted(IMAGE_MAP.keys()):
        img_info = IMAGE_MAP[img_filename]
        img_path = os.path.join(images_dir, img_filename)
        
        if not os.path.exists(img_path):
            print(f"⚠️  Skipping {img_filename} - file not found")
            continue
        
        section_type = img_info["section"]
        
        # Find appropriate insertion point
        if section_type == "Elara":
            keywords = ["elara", "nightshade", "protagonist"]
        elif section_type == "Lysandria":
            keywords = ["lysandria", "queen", "antagonist"]
        elif section_type == "Seraphiel":
            keywords = ["seraphiel", "celestial"]
        elif section_type == "locations":
            keywords = ["location", "setting", "world"]
        elif section_type == "scenes":
            keywords = ["scene", "key moment", "climax"]
        else:
            keywords = []
        
        para_index = find_paragraph_by_keyword(doc, keywords)
        
        if para_index is not None:
            # Add some space before the image
            doc.paragraphs[para_index]._element.addnext(
                doc.paragraphs[para_index]._element
            )
            
            print(f"✓ Adding {img_filename} near relevant section")
        else:
            print(f"→ Adding {img_filename} at end of document")
        
        # For simplicity, let's just add at the end with clear section headers
        added_images.append((img_path, img_info))
    
    return added_images


def add_images_to_document(docx_path, images_dir, output_path=None):
    """Main function to add images to the document"""
    
    if not os.path.exists(docx_path):
        print(f"ERROR: Document not found: {docx_path}")
        return False
    
    if not os.path.exists(images_dir):
        print(f"ERROR: Images directory not found: {images_dir}")
        return False
    
    print(f"📄 Loading document: {docx_path}")
    doc = Document(docx_path)
    
    print(f"📁 Images directory: {images_dir}")
    print(f"🖼️  Processing {len(IMAGE_MAP)} images...\n")
    
    # Strategy: Add images at the end with clear section breaks
    # This is safer than trying to insert in the middle
    
    # Add a page break and section header
    doc.add_page_break()
    
    # Add main section header
    header = doc.add_heading('Visual Reference Gallery', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Character Portraits Section
    doc.add_heading('Character Portraits', level=2)
    
    for img_filename in sorted(IMAGE_MAP.keys()):
        if IMAGE_MAP[img_filename]["section"] in ["Elara", "Lysandria", "Seraphiel"]:
            img_info = IMAGE_MAP[img_filename]
            img_path = os.path.join(images_dir, img_filename)
            
            if os.path.exists(img_path):
                # Add title
                title_para = doc.add_paragraph()
                title_para.add_run(img_info["title"]).bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(img_path, width=Inches(img_info["width"]))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add spacing
                doc.add_paragraph()
                
                print(f"✓ Added: {img_filename}")
            else:
                print(f"⚠️  Missing: {img_filename}")
    
    # Locations Section
    doc.add_page_break()
    doc.add_heading('Key Locations', level=2)
    
    for img_filename in sorted(IMAGE_MAP.keys()):
        if IMAGE_MAP[img_filename]["section"] == "locations":
            img_info = IMAGE_MAP[img_filename]
            img_path = os.path.join(images_dir, img_filename)
            
            if os.path.exists(img_path):
                # Add title
                title_para = doc.add_paragraph()
                title_para.add_run(img_info["title"]).bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(img_path, width=Inches(img_info["width"]))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add spacing
                doc.add_paragraph()
                
                print(f"✓ Added: {img_filename}")
            else:
                print(f"⚠️  Missing: {img_filename}")
    
    # Scenes Section
    doc.add_page_break()
    doc.add_heading('Key Scenes', level=2)
    
    for img_filename in sorted(IMAGE_MAP.keys()):
        if IMAGE_MAP[img_filename]["section"] == "scenes":
            img_info = IMAGE_MAP[img_filename]
            img_path = os.path.join(images_dir, img_filename)
            
            if os.path.exists(img_path):
                # Add title
                title_para = doc.add_paragraph()
                title_para.add_run(img_info["title"]).bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(img_path, width=Inches(img_info["width"]))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add spacing
                doc.add_paragraph()
                
                print(f"✓ Added: {img_filename}")
            else:
                print(f"⚠️  Missing: {img_filename}")
    
    # Save the document
    if output_path is None:
        # Create backup and overwrite original
        backup_path = docx_path.replace('.docx', '_backup.docx')
        os.replace(docx_path, backup_path)
        output_path = docx_path
        print(f"\n💾 Backup saved: {backup_path}")
    
    doc.save(output_path)
    print(f"✅ Document saved: {output_path}")
    
    return True


if __name__ == "__main__":
    # Set up paths
    script_dir = Path(__file__).parent
    project_dir = script_dir / "2025_10_19_Claude_Skills_Practice"
    
    # Use current directory if project structure not found
    if not project_dir.exists():
        project_dir = Path.cwd()
    
    docx_file = project_dir / "Blood_Assassin_Character_Dossier.docx"
    images_dir = project_dir / "generated-images"
    
    # Allow command line override
    if len(sys.argv) > 1:
        docx_file = Path(sys.argv[1])
    if len(sys.argv) > 2:
        images_dir = Path(sys.argv[2])
    
    print("=" * 60)
    print("Blood Assassin Character Dossier - Image Insertion")
    print("=" * 60)
    print()
    
    success = add_images_to_document(str(docx_file), str(images_dir))
    
    if success:
        print("\n" + "=" * 60)
        print("✅ SUCCESS! Images added to document.")
        print("=" * 60)
        sys.exit(0)
    else:
        print("\n" + "=" * 60)
        print("❌ FAILED. Check errors above.")
        print("=" * 60)
        sys.exit(1)
